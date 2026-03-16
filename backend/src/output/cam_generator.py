"""
cam_generator.py
────────────────
Generates a professional Credit Appraisal Memo (CAM) as a Word .docx file.
Formatted to Indian banking standards.

Document structure:
  1. Cover Page
  2. Executive Summary
  3. Company Background
  4. Financial Analysis
  5. Five Cs Scorecard
  6. Risk Flag Register
  7. Recommendation
  8. Conditions Precedent
"""

import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────
# STYLE HELPERS
# ─────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _bold_run(para, text: str, size: int = 11, color: str = None):
    run       = para.add_run(text)
    run.bold  = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run

def _normal_run(para, text: str, size: int = 10):
    run           = para.add_run(text)
    run.font.size = Pt(size)
    return run

def _add_heading(doc, text: str, level: int = 1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)
    run       = para.add_run(text.upper())
    run.bold  = True
    run.font.size = Pt(11 if level == 1 else 10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Add bottom border
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    pPr.append(pBdr)
    return para

def _add_kv_table(doc, rows: list[tuple], col_widths=(5, 10)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        kc = table.cell(i, 0)
        vc = table.cell(i, 1)
        _set_cell_bg(kc, "F3F3F3")
        kp     = kc.paragraphs[0]
        kp.paragraph_format.space_before = Pt(3)
        kp.paragraph_format.space_after  = Pt(3)
        _bold_run(kp, str(key), size=9)
        vp     = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after  = Pt(3)
        _normal_run(vp, str(val), size=9)
    return table


# ─────────────────────────────────────────────────────────────
# SECTION BUILDERS
# ─────────────────────────────────────────────────────────────

def _build_cover(doc, data: dict):
    # Institution header
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bold_run(hdr, "CREDIT APPRAISAL MEMORANDUM", size=16)
    hdr.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _normal_run(sub, "CONFIDENTIAL — FOR INTERNAL USE ONLY", size=9)

    doc.add_paragraph()

    # Company details box
    table = doc.add_table(rows=1, cols=1)
    cell  = table.cell(0, 0)
    _set_cell_bg(cell, "000000")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(data.get("company_name", "").upper())
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # Meta details
    scoring = data.get("scoring", {})
    _add_kv_table(doc, [
        ("Assessment ID",        data.get("analysis_id", "")),
        ("Date",                 data.get("date", datetime.now().strftime("%d %B %Y"))),
        ("CIN",                  data.get("cin", "")),
        ("Facility Type",        "Working Capital Facility"),
        ("Amount Requested",     f"₹{data.get('loan_amount_cr', 0)} Crore"),
        ("Industry Sector",      data.get("sector", "")),
        ("Final Score",          f"{scoring.get('composite_score', 0)} / 100"),
        ("Decision",             scoring.get("decision", "")),
    ])

    doc.add_page_break()


def _build_executive_summary(doc, data: dict):
    _add_heading(doc, "1. Executive Summary")
    scoring  = data.get("scoring", {})
    decision = scoring.get("decision", "PENDING")
    rationale = scoring.get("rationale", "")
    score    = scoring.get("composite_score", 0)

    # Decision callout
    table = doc.add_table(rows=1, cols=1)
    cell  = table.cell(0, 0)
    color = "1B5E20" if decision == "APPROVED" else "B71C1C" if decision == "REJECTED" else "E65100"
    _set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(f"DECISION: {decision}  |  COMPOSITE SCORE: {score}/100")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # Rationale
    para = doc.add_paragraph()
    _normal_run(para, rationale, size=10)

    # Cross-check summary
    cc = data.get("cross_checks", {})
    if any(cc.values()):
        doc.add_paragraph()
        para2 = doc.add_paragraph()
        _bold_run(para2, "Key Quantitative Findings:  ", size=10)
        findings = []
        if cc.get("circular_trading_instances", 0) > 0:
            findings.append(f"Circular trading: {cc['circular_trading_instances']} instances (₹{cc.get('circular_trading_cr',0):.2f}Cr)")
        if cc.get("cheque_bounces", 0) > 0:
            findings.append(f"Cheque bounces: {cc['cheque_bounces']} in statement period")
        if cc.get("hidden_emis", 0) > 0:
            findings.append(f"Hidden EMI patterns: {cc['hidden_emis']} detected")
        _normal_run(para2, " | ".join(findings), size=10)


def _build_five_cs(doc, data: dict):
    _add_heading(doc, "2. Five Cs Scorecard")
    scoring  = data.get("scoring", {})
    pillars  = scoring.get("pillars", {})

    # Summary table
    summary_rows = [("Pillar", "Weight", "Score", "Contribution", "Primary Driver")]
    for name, pillar in pillars.items():
        lines = pillar.get("score_lines", [])
        driver = lines[0].get("data_point", "") if lines else ""
        if len(driver) > 60:
            driver = driver[:57] + "..."
        summary_rows.append((
            name,
            f"{pillar.get('weight_pct')}%",
            f"{pillar.get('raw_score'):.1f} / 10",
            f"{pillar.get('weighted_contribution'):.1f} pts",
            driver,
        ))

    table = doc.add_table(rows=len(summary_rows), cols=5)
    table.style = "Table Grid"

    for i, row_data in enumerate(summary_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            p    = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            if i == 0:
                _set_cell_bg(cell, "000000")
                r = p.add_run(str(cell_text))
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                _normal_run(p, str(cell_text), size=9)
                if j == 2:  # score column — colour by value
                    try:
                        score_val = float(str(cell_text).split("/")[0].strip())
                        bg = "FFEBEE" if score_val < 4 else "FFF8E1" if score_val < 6 else "E8F5E9"
                        _set_cell_bg(cell, bg.replace("#", ""))
                    except:
                        pass

    # Detailed breakdown per pillar
    doc.add_paragraph()
    for name, pillar in pillars.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        _bold_run(p, f"{name.upper()}  ", size=10)
        _normal_run(p, f"(Weight: {pillar.get('weight_pct')}%  |  Score: {pillar.get('raw_score'):.1f}/10  |  Contribution: {pillar.get('weighted_contribution'):.1f} pts)", size=9)

        for line in pillar.get("score_lines", []):
            lp = doc.add_paragraph(style="List Bullet")
            lp.paragraph_format.left_indent = Cm(0.5)
            _bold_run(lp, f"{line.get('sub_factor', '')}: ", size=9)
            _normal_run(lp, f"{line.get('data_point', '')}  ", size=9)
            _normal_run(lp, f"[Adj: {line.get('adjustment', 0):+.1f}]  ", size=8)
            r = lp.add_run(f"Source: {line.get('source', '')}")
            r.font.size = Pt(8)
            r.italic = True


def _build_risk_flags(doc, data: dict):
    _add_heading(doc, "3. Risk Flag Register")

    research = data.get("research", {})
    flags    = research.get("risk_flags", [])

    if not flags:
        p = doc.add_paragraph()
        _normal_run(p, "No significant risk flags identified through research.", size=10)
        return

    # Sort by severity
    order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    flags = sorted(flags, key=lambda f: order.get(f.get("severity", "LOW"), 3))

    header_row = ["#", "Severity", "Type", "Finding", "Source", "Score Impact"]
    table = doc.add_table(rows=1 + len(flags), cols=6)
    table.style = "Table Grid"

    # Header
    for j, h in enumerate(header_row):
        cell = table.cell(0, j)
        _set_cell_bg(cell, "000000")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    # Rows
    for i, flag in enumerate(flags):
        sev   = flag.get("severity", "LOW")
        bg    = "FFEBEE" if sev == "HIGH" else "FFF9C4" if sev == "MEDIUM" else "F5F5F5"
        cells = [
            str(i + 1),
            sev,
            flag.get("type", ""),
            flag.get("title", ""),
            flag.get("source", ""),
            f"{float(flag.get('score_impact', 0)):+.1f} pts",
        ]
        for j, text in enumerate(cells):
            cell = table.cell(i + 1, j)
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(str(text))
            r.font.size = Pt(8)
            if j == 1 and sev == "HIGH":
                r.bold = True
                r.font.color.rgb = RGBColor(183, 28, 28)


def _build_recommendation(doc, data: dict):
    _add_heading(doc, "4. Recommendation")
    scoring  = data.get("scoring", {})
    decision = scoring.get("decision", "PENDING")
    rationale = scoring.get("rationale", "")
    limit    = scoring.get("recommended_limit_cr", 0)
    rate     = scoring.get("rate_premium_pct", 0)
    conditions = scoring.get("conditions_precedent", [])

    _add_kv_table(doc, [
        ("Decision",                  decision),
        ("Recommended Facility Limit", f"₹{limit} Crore"),
        ("Rate of Interest",          f"Base Rate + {rate}%" if rate else "N/A"),
        ("Composite Score",           f"{scoring.get('composite_score', 0)} / 100"),
    ])

    doc.add_paragraph()
    para = doc.add_paragraph()
    _bold_run(para, "Rationale:  ", size=10)
    _normal_run(para, rationale, size=10)

    if conditions:
        doc.add_paragraph()
        _add_heading(doc, "5. Conditions Precedent", level=2)
        for c in conditions:
            cp = doc.add_paragraph(style="List Bullet")
            _normal_run(cp, c, size=10)


def _build_footer_disclaimer(doc):
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell  = table.cell(0, 0)
    _set_cell_bg(cell, "F5F5F5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(
        "This memorandum was generated with AI assistance by Intelli-Credit v1.0. "
        "It must be reviewed and countersigned by an authorized credit officer "
        "before any lending decision is made."
    )
    r.font.size = Pt(8)
    r.italic    = True


# ─────────────────────────────────────────────────────────────
# MASTER GENERATOR
# ─────────────────────────────────────────────────────────────

def generate_cam(analysis_data: dict, qualitative_data: dict = None) -> str:
    """
    Generate a Word CAM document from analysis data.
    Returns the file path of the generated .docx file.
    """
    print(f"\n  Generating CAM for {analysis_data.get('company_name')}...")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Build sections
    _build_cover(doc, analysis_data)
    _build_executive_summary(doc, analysis_data)
    doc.add_paragraph()
    _build_five_cs(doc, analysis_data)
    doc.add_paragraph()
    _build_risk_flags(doc, analysis_data)
    doc.add_paragraph()
    _build_recommendation(doc, analysis_data)
    _build_footer_disclaimer(doc)

    # Save
    output_dir  = tempfile.gettempdir()
    safe_name   = analysis_data.get("company_name", "Company").replace(" ", "_")
    analysis_id = analysis_data.get("analysis_id", "CAM")
    output_path = os.path.join(output_dir, f"CAM_{safe_name}_{analysis_id}.docx")
    doc.save(output_path)

    print(f"  CAM saved: {output_path}")
    return output_path